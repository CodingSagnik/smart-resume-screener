import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract raw text from a PDF file using pdfplumber with fallback to pypdf.
    Handles encrypted, malformed, or image-only PDFs gracefully.
    """
    text_content = []

    # 1. Primary Strategy: pdfplumber (superior layout, table, and multi-column parsing)
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            if len(pdf.pages) == 0:
                raise ValueError("PDF file has 0 pages.")

            for page_idx, page in enumerate(pdf.pages):
                page_text = page.extract_text(layout=True) or page.extract_text()
                if page_text and page_text.strip():
                    text_content.append(page_text.strip())

        extracted = "\n\n".join(text_content).strip()
        if extracted:
            return extracted
        logger.warning("pdfplumber extracted empty text. Attempting fallback with pypdf...")
    except ImportError:
        logger.info("pdfplumber not installed, attempting fallback to pypdf.")
    except Exception as e:
        logger.warning(f"pdfplumber failed with error: {e}. Attempting fallback with pypdf...")

    # 2. Fallback Strategy: pypdf
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                raise ValueError("PDF is encrypted and password-protected.")

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_content.append(page_text.strip())

        extracted = "\n\n".join(text_content).strip()
        if extracted:
            return extracted
    except Exception as e:
        logger.error(f"pypdf extraction failed: {e}")
        raise ValueError(f"Failed to read PDF file: {str(e)}")

    # 3. Check if document was an unsearchable image scan
    if not text_content:
        raise ValueError(
            "No readable text could be extracted from the PDF. "
            "The document might be an image scan or contain no text elements."
        )

    return "\n\n".join(text_content).strip()


def extract_text_from_txt(file_path: str) -> str:
    """Extract raw text from a plain text file with multiple encoding fallbacks."""
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "iso-8859-1"]
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read().strip()
                if text:
                    return text
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception as e:
            raise ValueError(f"Failed to read text file: {str(e)}")

    raise ValueError("Could not decode text file with standard encodings.")


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw text from a Microsoft Word (.docx) document."""
    try:
        import docx

        doc = docx.Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        extracted = "\n".join(full_text).strip()
        if not extracted:
            raise ValueError("DOCX document is empty.")
        return extracted
    except ImportError:
        raise ValueError("python-docx library is not installed for docx parsing.")
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_file(file_path: str, file_type: Optional[str] = None) -> str:
    """
    Main dispatcher for text extraction from various resume file formats.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at: {file_path}")

    if not file_type:
        file_type = file_path.split(".")[-1].lower() if "." in file_path else ""

    file_type = file_type.lower().strip(".")

    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type in ["txt", "text", "log"]:
        return extract_text_from_txt(file_path)
    elif file_type in ["docx", "doc"]:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: .{file_type}")
